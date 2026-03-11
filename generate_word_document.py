import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
import requests
from io import BytesIO
import re
import urllib3

# Suppress SSL warnings (if verify=False is used)
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# Load the CSV
file_path = "issues.csv"
df = pd.read_csv(file_path)

# Clean only string columns to avoid pandas warnings
df[df.select_dtypes(include='object').columns] = df.select_dtypes(include='object').fillna('')

# Success Criteria mapping
success_criteria_mapping = {
    "1.1.1": "1.1.1 Non-Text Content",
    "1.2.1": "1.2.1 Audio-Only and Video-Only",
    "1.2.2": "1.2.2 Captions (Prerecorded)",
    "1.2.3": "1.2.3 Audio Description or Media Alternative",
    "1.2.4": "1.2.4 Captions (Live)",
    "1.2.5": "1.2.5 Audio Description (Pre-Recorded)",
    "1.3.1": "1.3.1 Info and Relationships",
    "1.3.2": "1.3.2 Meaningful Sequence",
    "1.3.3": "1.3.3 Sensory Characteristics",
    "1.3.4": "1.3.4 Orientation",
    "1.3.5": "1.3.5 Identify Input Purpose",
    "1.4.1": "1.4.1 Use of Color",
    "1.4.2": "1.4.2 Audio Control",
    "1.4.3": "1.4.3 Contrast Minimum",
    "1.4.4": "1.4.4 Resize Text",
    "1.4.5": "1.4.5 Images of Text",
    "1.4.10": "1.4.10 Reflow",
    "1.4.11": "1.4.11 Non-Text Contrast",
    "1.4.12": "1.4.12 Text Spacing",
    "1.4.13": "1.4.13 Content on Hover or Focus",
    "2.1.1": "2.1.1 Keyboard",
    "2.1.2": "2.1.2 No Keyboard Trap",
    "2.1.4": "2.1.4 Character Key Shortcuts",
    "2.2.1": "2.2.1 Timing Adjustable",
    "2.2.2": "2.2.2 Pause, Stop, Hide",
    "2.3.1": "2.3.1 Three Flashes or Below",
    "2.4.1": "2.4.1 Bypass Blocks",
    "2.4.2": "2.4.2 Page Titled",
    "2.4.3": "2.4.3 Focus Order",
    "2.4.4": "2.4.4 Link Purpose (In Context)",
    "2.4.5": "2.4.5 Multiple Ways",
    "2.4.6": "2.4.6 Headings and Labels",
    "2.4.7": "2.4.7 Focus Visible",
    "2.4.11": "2.4.11 Focus Not Obscured",
    "2.5.1": "2.5.1 Pointer Gestures",
    "2.5.2": "2.5.2 Pointer Cancellation",
    "2.5.3": "2.5.3 Label in Name",
    "2.5.4": "2.5.4 Motion Actuation",
    "2.5.5": "2.5.5 Target Size",
    "2.5.7": "2.5.7 Dragging Movements",
    "2.5.8": "2.5.8 Target Size",
    "3.1.1": "3.1.1 Language of Page",
    "3.1.2": "3.1.2 Language of Parts",
    "3.2.1": "3.2.1 On Focus",
    "3.2.2": "3.2.2 On Input",
    "3.2.3": "3.2.3 Consistent Navigation",
    "3.2.4": "3.2.4 Consistent Identification",
    "3.2.6": "3.2.6 Consistent Help",
    "3.3.1": "3.3.1 Error Identification",
    "3.3.2": "3.3.2 Labels or Instruction",
    "3.3.3": "3.3.3 Error Suggestion",
    "3.3.4": "3.3.4 Error Prevention (Legal, Financial, Data)",
    "3.3.7": "3.3.7 Redundant Entry",
    "3.3.8": "3.3.8 Accessible Authentication",
    "4.1.2": "4.1.2 Name, Role, Value",
    "4.1.3": "4.1.3 Status Messages"
}

# Create Word document
doc = Document()
doc.add_heading("Accessibility Issues Report", level=1)

# URL regex pattern
url_pattern = r'(https?://[^\s,]+)'

# Grouping by Success Criteria and Test Unit
for success_criteria, sc_group in df.groupby('Success Criteria'):
    # Styled Heading Level 3 (Success Criteria)
    mapped_text = success_criteria_mapping.get(success_criteria.strip(), success_criteria)
    sc_heading = doc.add_heading(level=3)
    sc_run = sc_heading.add_run(mapped_text)
    sc_run.font.name = 'Times New Roman'
    sc_run.font.color.rgb = RGBColor(0, 78, 154)
    sc_run.font.size = Pt(16)
    sc_run.font.bold = True
    sc_run.font.italic = False
    sc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for test_unit, tu_group in sc_group.groupby('Test Unit'):
        # Styled Heading Level 4 (Test Unit)
        tu_heading = doc.add_heading(level=4)
        tu_run = tu_heading.add_run(test_unit)
        tu_run.font.name = 'Times New Roman'
        tu_run.font.color.rgb = RGBColor(0, 0, 0)
        tu_run.font.size = Pt(15)
        tu_run.font.bold = True
        tu_run.font.italic = False
        tu_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        for _, row in tu_group.iterrows():
            # Styled Heading Level 5 (Summary)
            summary_heading = doc.add_heading(level=5)
            summary_run = summary_heading.add_run(row['Summary'])
            summary_run.font.name = 'Times New Roman'
            summary_run.font.color.rgb = RGBColor(0, 0, 0)
            summary_run.font.size = Pt(14)
            summary_run.font.bold = True
            summary_run.font.italic = False
            summary_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Styled Heading Level 6 - Description
            desc_heading = doc.add_heading(level=6)
            desc_run = desc_heading.add_run("Description")
            desc_run.font.name = 'Times New Roman'
            desc_run.font.color.rgb = RGBColor(0, 0, 0)
            desc_run.font.size = Pt(14)
            desc_run.font.bold = True
            desc_run.font.italic = False
            desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            doc.add_paragraph(row['Description'])

            # Styled Heading Level 6 - Impact inline
            impact_text = f"Impact: {row['Impact']}"
            impact_heading = doc.add_heading(level=6)
            impact_run = impact_heading.add_run(impact_text)
            impact_run.font.name = 'Times New Roman'
            impact_run.font.color.rgb = RGBColor(0, 0, 0)
            impact_run.font.size = Pt(14)
            impact_run.font.bold = True
            impact_run.font.italic = False
            impact_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

                # Styled Heading Level 6 - HTML Code
            html_heading = doc.add_heading(level=6)
            html_run = html_heading.add_run("HTML Code")
            html_run.font.name = 'Times New Roman'
            html_run.font.color.rgb = RGBColor(0, 0, 0)
            html_run.font.size = Pt(14)
            html_run.font.bold = True
            html_run.font.italic = False
            html_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            doc.add_paragraph(row['Source Code'])

            # Styled Heading Level 6 - Screenshots
            screen_heading = doc.add_heading(level=6)
            screen_run = screen_heading.add_run("Screenshots")
            screen_run.font.name = 'Times New Roman'
            screen_run.font.color.rgb = RGBColor(0, 0, 0)
            screen_run.font.size = Pt(14)
            screen_run.font.bold = True
            screen_run.font.italic = False
            screen_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            screenshots_text = row['Screenshots']
            urls = re.findall(url_pattern, screenshots_text)

            if urls:
                for idx, url in enumerate(urls, start=1):
                    try:
                        response = requests.get(url, timeout=10, verify=False)
                        response.raise_for_status()
                        image_stream = BytesIO(response.content)
                        doc.add_picture(image_stream, width=Inches(5))
                        doc.paragraphs[-1].add_run(f"\n(Screenshot {idx})")
                    except Exception as e:
                        doc.add_paragraph(f"[Error downloading image from {url}]: {e}")
            else:
                doc.add_paragraph("[No valid screenshot URL found]")

            # Styled Heading Level 6 - Recommended to fix
            fix_heading = doc.add_heading(level=6)
            fix_run = fix_heading.add_run("Recommended to fix")
            fix_run.font.name = 'Times New Roman'
            fix_run.font.color.rgb = RGBColor(0, 0, 0)
            fix_run.font.size = Pt(14)
            fix_run.font.bold = True
            fix_run.font.italic = False
            fix_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            doc.add_paragraph(row['Recommended to fix'])

# Save the final document
output_path = "Accessibility_Issues_Report.docx"
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")